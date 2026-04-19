"""Regression tests for the publication-quality figure toolkit."""

from __future__ import annotations

import shutil
from contextlib import contextmanager
from pathlib import Path
from uuid import uuid4

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

from fintools.datasets import available_validation_datasets, load_validation_dataset
from fintools.figures import (
    FT_BACKGROUND,
    FigureContext,
    WordFigureEntry,
    create_figure_suite,
    cumulative_returns_plot,
    export_figure_bundle,
    export_word_figure,
    insert_figures_docx,
    lollipop_plot,
    plan_figure_suite,
    profile_dataframe,
    theme_rc,
    validate_docx_images_fit_page,
    validate_image_not_blank,
)


ROOT = Path(__file__).resolve().parents[1]
TMP_ROOT = ROOT / ".tmp-figure-tests"


def close(fig) -> None:
    plt.close(fig)


def docx_text(path: Path) -> str:
    """Return text from top-level paragraphs and caption tables."""

    from docx import Document

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


@contextmanager
def temp_figure_dir():
    """Create a repo-local scratch figure directory."""

    TMP_ROOT.mkdir(exist_ok=True)
    path = TMP_ROOT / f"figures-{uuid4().hex}"
    path.mkdir(parents=True)
    try:
        yield path
    finally:
        shutil.rmtree(path, ignore_errors=True)


def test_validation_datasets_load_with_metadata() -> None:
    expected = {
        "ff3_monthly",
        "ff25_size_value_monthly",
        "ff_industry_10_monthly",
        "fred_financial_stress_daily",
        "fred_macro_monthly",
        "fred_rates_daily",
        "shiller_market_monthly",
        "world_bank_country_panel_annual",
        "world_bank_gdp_annual",
    }
    assert expected.issubset(set(available_validation_datasets()))

    ff3 = load_validation_dataset("ff3_monthly")
    assert isinstance(ff3.data.index, pd.DatetimeIndex)
    assert {"Mkt-RF", "SMB", "HML", "RF"}.issubset(ff3.data.columns)
    assert ff3.units["Mkt-RF"] == "percent monthly return"
    assert "Kenneth French" in ff3.source

    gdp = load_validation_dataset("world_bank_gdp_annual")
    assert {"country", "country_code", "gdp_trillions_usd"}.issubset(gdp.data.columns)
    assert "World Bank" in gdp.source


def test_ft_theme_is_opt_in_and_background_is_explicit() -> None:
    fins_rc = theme_rc()
    ft_rc = theme_rc(style="ft")
    ft_background_rc = theme_rc(style="ft", ft_background=True)

    assert fins_rc["figure.facecolor"] == "white"
    assert ft_rc["figure.facecolor"] == "white"
    assert ft_background_rc["figure.facecolor"] == FT_BACKGROUND
    assert ft_rc["axes.prop_cycle"] != fins_rc["axes.prop_cycle"]


def test_export_bundle_and_word_docx() -> None:
    ff3 = load_validation_dataset("ff3_monthly").data
    context = FigureContext(
        title="Cumulative Market Excess Return",
        note="Validation figure for Word/A4 export.",
        source="Kenneth French Data Library",
        sample="2020-01-31 to 2021-12-31",
        units="Cumulative return",
    )
    fig, _ = cumulative_returns_plot(
        ff3,
        "Mkt-RF",
        returns_are_percent=True,
        wealth_index=True,
        log_scale=True,
        profile="word_a4",
        style="ft",
    )

    with temp_figure_dir() as output_dir:
        bundle = export_figure_bundle(fig, output_dir, "market_cumulative", context=context)
        assert bundle["png"].exists()
        assert bundle["pdf"].exists()
        assert "Kenneth French" in bundle["caption"].read_text(encoding="utf-8")
        assert not validate_image_not_blank(bundle["png"])

        word_bundle = export_word_figure(
            fig,
            output_dir,
            "market_cumulative_word",
            context=context,
        )
        docx_path = insert_figures_docx(
            [WordFigureEntry(word_bundle["png"], context=context)],
            output_dir / "figure_pack.docx",
            title="Validation Figure Pack",
        )
        assert not validate_docx_images_fit_page(docx_path)
        text = docx_text(docx_path)
        assert "Validation Figure Pack" in text
        assert "Figure 1. Cumulative Market Excess Return." in text
    close(fig)


def test_dataframe_figure_suite_profiles_plans_and_exports() -> None:
    rng = np.random.default_rng(123)
    dates = pd.date_range("2018-01-31", periods=72, freq="ME")
    frame = pd.DataFrame(
        {
            "date": dates,
            "market_return": rng.normal(0.7, 3.8, len(dates)),
            "smb_return": rng.normal(0.1, 1.7, len(dates)),
            "credit_spread_percent": rng.normal(3.0, 0.35, len(dates)).clip(0.5),
            "volume_usd": np.linspace(100, 165, len(dates)) + rng.normal(0, 4, len(dates)),
            "segment": np.where(np.arange(len(dates)) % 2 == 0, "IG", "HY"),
        }
    )

    profile = profile_dataframe(frame)
    assert profile.date_column == "date"
    assert "segment" in profile.categorical_columns
    assert {"market_return", "smb_return"}.issubset(profile.return_columns)

    plan = plan_figure_suite(frame, title_prefix="Demo", max_figures=4)
    assert [item.kind for item in plan] == [
        "wide_time_series",
        "indexed_time_series",
        "cumulative_returns",
        "mean_bar",
    ]

    with temp_figure_dir() as output_dir:
        result = create_figure_suite(
            frame,
            output_dir,
            style="ft",
            docx=True,
            source="",
            title_prefix="Demo",
            max_figures=4,
            formats=("png",),
        )
        assert not result.issues
        assert not result.skipped
        assert [figure.plan_item.kind for figure in result.generated_figures] == [
            "wide_time_series",
            "indexed_time_series",
            "cumulative_returns",
            "mean_bar",
        ]
        assert result.docx_path is not None
        assert not validate_docx_images_fit_page(result.docx_path)


def test_lollipop_highlights_have_distinct_colors() -> None:
    gdp = load_validation_dataset("world_bank_gdp_annual").data.reset_index()
    latest = gdp[gdp["date"].dt.year == 2024].copy()
    fig, ax = lollipop_plot(
        latest,
        "country",
        "gdp_trillions_usd",
        highlight=["United States", "China", "India"],
        style="ft",
    )

    point_colors = {
        tuple(facecolor)
        for collection in ax.collections
        for facecolor in collection.get_facecolors()  # type: ignore[attr-defined]
    }
    assert len(point_colors) >= 4
    close(fig)


def test_figure_workflow_docs_and_ignore_policy() -> None:
    docs = (ROOT / "docs" / "ai" / "figures.md").read_text(encoding="utf-8")
    assert "tools/figure_examples.py --style ft --docx --output results/figures" in docs
    assert "fintools.figures" in docs
    assert 'style="fins"' in docs
    assert 'style="ft"' in docs

    context_files = [
        ROOT / "AGENTS.md",
        ROOT / "CLAUDE.md",
        ROOT / "GEMINI.md",
        ROOT / "docs" / "ai" / "writing.md",
        ROOT / ".claude" / "rules" / "latex-conventions.md",
        ROOT / ".claude" / "skills" / "new-project" / "SKILL.md",
        ROOT / ".claude" / "skills" / "publication-figures" / "SKILL.md",
    ]
    for path in context_files:
        text = path.read_text(encoding="utf-8")
        assert "docs/ai/figures.md" in text
        assert "fintools.figures" in text

    skill_text = (ROOT / ".claude" / "skills" / "publication-figures" / "SKILL.md").read_text(encoding="utf-8")
    assert "finance.mplstyle" in skill_text
    assert "figutils.py" in skill_text
    assert (ROOT / ".claude" / "skills" / "publication-figures" / "finance.mplstyle").exists()
    assert (ROOT / ".claude" / "skills" / "publication-figures" / "figutils.py").exists()

    gitignore = (ROOT / ".gitignore").read_text(encoding="utf-8")
    for expected in [
        ".maintainers/",
        "fintools/datasets/validation/_raw/",
        "fintools/datasets/validation/_refresh/",
        "**/results/figures/*.docx",
        "**/results/figures/*.png",
        "**/results/figures/*.pdf",
        "**/results/figures/*.caption.md",
    ]:
        assert expected in gitignore
